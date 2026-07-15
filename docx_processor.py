"""Read/write .docx files while preserving Word formatting.

All text replacement happens at the level of individual `run` objects, never
by reassigning `paragraph.text` wholesale, so that fonts, bold/italic, colors
and other direct (run-level) formatting survive the round trip.
"""

from __future__ import annotations

from typing import Dict, Iterable, List, Tuple

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

# ---------------------------------------------------------------------------
# Traversal helpers
# ---------------------------------------------------------------------------


def iter_table_paragraphs(table: Table) -> Iterable[Paragraph]:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested_table in cell.tables:
                yield from iter_table_paragraphs(nested_table)


def iter_all_paragraphs(document: Document) -> Iterable[Paragraph]:
    """Yield every paragraph in the document body, tables, headers and footers."""
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        yield from iter_table_paragraphs(table)
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for paragraph in section.footer.paragraphs:
            yield paragraph
        for table in section.header.tables:
            yield from iter_table_paragraphs(table)
        for table in section.footer.tables:
            yield from iter_table_paragraphs(table)


def extract_docx_text(document: Document) -> str:
    """Best-effort plain-text extraction (paragraphs + tables) for preview."""
    parts = [p.text for p in iter_all_paragraphs(document) if p.text.strip()]
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Run-level replacement
# ---------------------------------------------------------------------------


def apply_replacements_to_paragraph(
    paragraph: Paragraph, matches: List[Tuple[int, int, str]]
) -> None:
    """Replace substrings of the paragraph's text with `matches`.

    `matches` is a list of (start, end, replacement) offsets relative to the
    paragraph's *original* text (before any replacement), sorted ascending
    and non-overlapping. A match may span multiple runs (Word frequently
    splits a single word across runs, e.g. due to spell-check markers); in
    that case the replacement text is inserted using the formatting of the
    run in which the match starts, and the remaining overlapped runs simply
    lose the overlapped portion of their text.
    """
    if not matches:
        return
    runs = list(paragraph.runs)
    if not runs:
        return

    spans = []
    pos = 0
    for run in runs:
        spans.append((pos, pos + len(run.text)))
        pos += len(run.text)

    new_texts = []
    for run, (run_start, run_end) in zip(runs, spans):
        pieces = []
        cursor = run_start
        for start, end, replacement in matches:
            ov_start, ov_end = max(start, run_start), min(end, run_end)
            if ov_start >= ov_end:
                continue
            pieces.append(run.text[cursor - run_start : ov_start - run_start])
            if start >= run_start:
                # This run contains the *start* of the match: insert the
                # replacement text exactly once, here.
                pieces.append(replacement)
            cursor = ov_end
        pieces.append(run.text[cursor - run_start :])
        new_texts.append("".join(pieces))

    for run, new_text in zip(runs, new_texts):
        run.text = new_text


def find_literal_matches(text: str, replacements: Dict[str, str]) -> List[Tuple[int, int, str]]:
    """Find every literal, non-overlapping occurrence of each key (token) in
    `replacements` inside `text`. Used for the reverse (restore) pass, where
    tokens like `[ФИО_1]` are located verbatim rather than via NLP offsets.
    """
    matches: List[Tuple[int, int, str]] = []
    for token, value in replacements.items():
        start = 0
        while True:
            idx = text.find(token, start)
            if idx == -1:
                break
            matches.append((idx, idx + len(token), value))
            start = idx + len(token)
    matches.sort(key=lambda m: m[0])
    return _drop_overlaps(matches)


def _drop_overlaps(matches: List[Tuple[int, int, str]]) -> List[Tuple[int, int, str]]:
    result = []
    last_end = -1
    for start, end, repl in matches:
        if start < last_end:
            continue
        result.append((start, end, repl))
        last_end = end
    return result
